# output_handlers.py
import os
import re
import pandas as pd
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import settings

def ensure_pdf_path():
    # 저장 경로가 비어 있으면 Save As로 받기
    if not getattr(settings, "output_file_pdf", ""):
        path = filedialog.asksaveasfilename(
            title="PDF 저장 위치 선택",
            defaultextension=".pdf",
            filetypes=[("PDF 파일", "*.pdf")]
        )
        if not path:
            return None
        settings.output_file_pdf = path
    return settings.output_file_pdf

# ===== 공용 유틸 =====
def open_file(file_path):
    try:
        os.startfile(file_path)
    except Exception as e:
        messagebox.showerror("파일 열기 오류", f"파일을 열 수 없습니다: {e}")

# ===== 엑셀 / DOCX / PDF 경로 세팅 및 저장 =====
def save_output_files(path_dir, file_type="all"):
    # 엑셀
    if file_type in ["all", "excel"]:
        df = pd.DataFrame(settings.data)
        settings.output_file_excel = os.path.join(path_dir, 'output.xlsx')

        split_data = []
        for entry in settings.data:
            content = entry['내용'].splitlines()
            split_row = [line.split() for line in content]
            split_data.extend(split_row)

        if split_data:
            max_cols = max(len(r) for r in split_data)
            for r in split_data:
                r.extend([''] * (max_cols - len(r)))
            df_split = pd.DataFrame(split_data)
        else:
            df_split = pd.DataFrame()
        df_split.to_excel(settings.output_file_excel, index=False, header=False)

    # DOCX
    if file_type in ["all", "docx"]:
        df = pd.DataFrame(settings.data)
        settings.output_file_docx = os.path.join(path_dir, 'output.docx')
        doc = Document()
        doc.add_heading('제목을 적어주세요.', 0)
        for _, row in df.iterrows():
            doc.add_paragraph(row['내용'])
        doc.save(settings.output_file_docx)

    # PDF (경로만 설정, 실 저장은 미리보기에서)
    if file_type in ["all", "pdf"]:
        _ = pd.DataFrame(settings.data)
        settings.output_file_pdf = os.path.join(path_dir, 'output.pdf')

    print(f"{file_type} 파일 저장 완료")

def add_text_to_existing_file(file_type, df):
    if file_type == "엑셀":
        file_path = filedialog.askopenfilename(
            title="엑셀 파일 선택",
            filetypes=[("엑셀 파일", "*.xlsx;*.xls")]
        )
    elif file_type == "DOCX":
        file_path = filedialog.askopenfilename(
            title="DOCX 파일 선택",
            filetypes=[("Word 파일", "*.docx")]
        )
    else:
        return

    if not file_path:
        return

    if file_type == "엑셀":
        try:
            existing_df = pd.read_excel(file_path, header=None, engine='openpyxl')
            new_data = df[['내용']]
            split_data = []
            for entry in new_data['내용']:
                content = entry.splitlines()
                for line in content:
                    split_row = line.split()
                    split_data.append(split_row)

            if split_data:
                max_cols = max(len(r) for r in split_data)
                for r in split_data:
                    r.extend([''] * (max_cols - len(r)))
                combined_df = pd.concat([existing_df, pd.DataFrame(split_data)], ignore_index=True)
            else:
                combined_df = existing_df

            combined_df.to_excel(file_path, index=False, header=False)
            messagebox.showinfo("성공", "엑셀 파일에 텍스트가 추가되었습니다.")
            open_file(file_path)
        except Exception as e:
            messagebox.showerror("오류", f"엑셀 파일 처리 중 오류가 발생했습니다: {e}")
            print("오류 발생:", e)

    elif file_type == "DOCX":
        try:
            existing_doc = Document(file_path)
            for _, row in df.iterrows():
                existing_doc.add_paragraph(row['내용'])
            existing_doc.save(file_path)
            messagebox.showinfo("성공", "DOCX 파일에 텍스트가 추가되었습니다.")
            open_file(file_path)
        except Exception as e:
            messagebox.showerror("오류", f"DOCX 파일 처리 중 오류가 발생했습니다: {e}")
            print("오류 발생:", e)

# ===== PDF 저장(커스텀 폰트/폰트사이즈 태그/이미지 삽입) =====
def save_to_pdf_with_custom_fonts(text_widget, font_map, pdf_path):
    # 폰트 파일 존재/등록 확인
    try:
        if not os.path.isfile(settings.font_path):
            messagebox.showerror("폰트 오류", f"폰트 파일을 찾을 수 없어요:\n{settings.font_path}")
            return
        pdfmetrics.registerFont(TTFont('NanumGothic', settings.font_path))
    except Exception as e:
        import traceback
        traceback.print_exc()
        messagebox.showerror("폰트 등록 실패", f"NanumGothic 등록 중 오류:\n{e}")
        return

    c = pdfcanvas.Canvas(pdf_path, pagesize=letter)
    pdfmetrics.registerFont(TTFont('NanumGothic', settings.font_path))
    width, height = letter
    margin = 50
    y_pos = height - margin

    index = "1.0"
    while True:
        if text_widget.compare(index, ">=", "end"):
            break

        line_text = text_widget.get(index, f"{index} +1line")

        # 이미지 토큰 파싱
        match = re.match(r"\[이미지:\s*(.*?),\s*(\d+),\s*(\d+)\]", line_text.strip())
        if match:
            image_path, w_str, h_str = match.groups()
            img_w, img_h = float(w_str), float(h_str)

            if y_pos - img_h < margin:
                c.showPage()
                y_pos = height - margin

            c.drawImage(image_path.strip(), margin, y_pos - img_h, width=img_w, height=img_h)
            y_pos -= img_h + 10
            index = text_widget.index(f"{index} +1line")
            continue

        if not line_text.strip():
            y_pos -= 20
            index = text_widget.index(f"{index} +1line")
            continue

        # 라인 내 최대 어센트/디센트 계산
        char_index = index
        max_ascent = 0
        max_descent = 0
        while True:
            ch = text_widget.get(char_index)
            if ch == "\n":
                break
            tags = text_widget.tag_names(char_index)
            font_size = 12
            for t in tags:
                if t.startswith("font_"):
                    try:
                        font_size = int(t.split("_")[-1])
                    except:
                        pass
            ascent = pdfmetrics.getAscent('NanumGothic') / 1000 * font_size
            descent = pdfmetrics.getDescent('NanumGothic') / 1000 * font_size
            max_ascent = max(max_ascent, ascent)
            max_descent = max(max_descent, descent)
            char_index = text_widget.index(f"{char_index} +1c")

        line_spacing = max_ascent + max_descent + 5
        if y_pos - line_spacing < margin:
            c.showPage()
            y_pos = height - margin

        baseline = y_pos - max_ascent

        # 실제 그리기(자동 줄바꿈)
        x_cursor = margin
        char_index = index
        while True:
            ch = text_widget.get(char_index)
            if ch == "\n":
                break
            tags = text_widget.tag_names(char_index)
            font_size = 12
            for t in tags:
                if t.startswith("font_"):
                    try:
                        font_size = int(t.split("_")[-1])
                    except:
                        pass

            ascent = pdfmetrics.getAscent('NanumGothic') / 1000 * font_size
            char_w = pdfmetrics.stringWidth(ch, 'NanumGothic', font_size)

            if x_cursor + char_w > width - margin:
                x_cursor = margin
                y_pos -= line_spacing
                baseline = y_pos - max_ascent
                if y_pos - line_spacing < margin:
                    c.showPage()
                    y_pos = height - margin
                    baseline = y_pos - max_ascent

            c.setFont('NanumGothic', font_size)
            c.drawString(x_cursor, baseline, ch)

            x_cursor += char_w
            char_index = text_widget.index(f"{char_index} +1c")

        y_pos -= line_spacing
        index = text_widget.index(f"{index} +1line")

    c.save()

# ===== PDF 미리보기/편집 GUI =====
def preview_and_edit_pdf():
    root = tk.Toplevel()  # 메인 메뉴 위에 뜨도록
    root.title("PDF 미리보기 및 수정")
    root.geometry("600x650")

    text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=70, height=25)
    text_area.pack(pady=10, padx=10)

    extracted_text = "\n\n".join([entry['내용'] for entry in settings.data])
    text_area.insert(tk.END, extracted_text)

    # 검색 UI
    search_frame = tk.Frame(root)
    search_frame.pack(pady=5)

    tk.Label(search_frame, text="검색어:").pack(side=tk.LEFT)
    search_entry = tk.Entry(search_frame, width=20)
    search_entry.pack(side=tk.LEFT, padx=5)

    def search_text():
        text_area.tag_remove("highlight", "1.0", tk.END)
        keyword = search_entry.get()
        if not keyword:
            return
        start = "1.0"
        while True:
            start = text_area.search(keyword, start, stopindex=tk.END)
            if not start:
                break
            end = f"{start}+{len(keyword)}c"
            text_area.tag_add("highlight", start, end)
            start = end
        text_area.tag_config("highlight", background="yellow", foreground="black")

    tk.Button(search_frame, text="검색", command=search_text).pack(side=tk.LEFT)

    # 폰트 사이즈 태그 적용
    font_size_frame = tk.Frame(root)
    font_size_frame.pack()
    tk.Label(font_size_frame, text="폰트 크기:").pack(side=tk.LEFT)
    font_size_entry = tk.Entry(font_size_frame, width=5)
    font_size_entry.insert(0, "12")
    font_size_entry.pack(side=tk.LEFT)

    def apply_font_to_selection():
        try:
            font_size = int(font_size_entry.get())
            start = text_area.index(tk.SEL_FIRST)
            end = text_area.index(tk.SEL_LAST)
            settings.font_size_map.append((start, end, font_size))
            tag_name = f"font_{font_size}"
            text_area.tag_add(tag_name, start, end)
            text_area.tag_config(tag_name, font=("NanumGothic", font_size))
        except tk.TclError:
            messagebox.showerror("오류", "텍스트를 먼저 선택해주세요.")
        except ValueError:
            messagebox.showerror("입력 오류", "숫자 형식의 폰트 크기를 입력해주세요.")

    tk.Button(root, text="선택 텍스트 크기 적용", command=apply_font_to_selection).pack(pady=5)

    # 이미지 삽입
    def add_image_to_pdf():
        file_path = filedialog.askopenfilename(title="이미지 파일 선택", filetypes=[("이미지 파일", "*.png;*.jpg;*.jpeg")])
        if not file_path:
            return
        size_window = tk.Toplevel(root)
        size_window.title("이미지 크기 조정")

        tk.Label(size_window, text="너비:").pack()
        width_entry = tk.Entry(size_window)
        width_entry.pack()

        tk.Label(size_window, text="높이:").pack()
        height_entry = tk.Entry(size_window)
        height_entry.pack()

        def insert_image():
            w = width_entry.get()
            h = height_entry.get()
            text_area.insert(tk.INSERT, f"\n[이미지: {file_path}, {w}, {h}]\n")
            size_window.destroy()

        tk.Button(size_window, text="추가", command=insert_image).pack()

    def save_edited_pdf():
        # 1) 폰트 사이즈 숫자 확인
        try:
            _ = int(font_size_entry.get())
        except ValueError:
            messagebox.showerror("입력 오류", "올바른 숫자 형식의 폰트 크기를 입력해주세요.")
            return

        # 2) PDF 저장 경로 보장(미리보기 버튼 안 거치고 온 경우 대비)
        pdf_path = ensure_pdf_path()
        if not pdf_path:
            messagebox.showinfo("취소됨", "PDF 저장이 취소되었습니다.")
            return

        # 디버그: 현재 경로 보여주기(문제 추적용)
        # messagebox.showinfo("디버그", f"저장 경로: {pdf_path}")

        # 3) 실제 저장 시도(예외는 팝업으로)
        try:
            save_to_pdf_with_custom_fonts(text_area, settings.font_size_map, pdf_path)
        except FileNotFoundError as e:
            messagebox.showerror("파일 오류", f"경로를 찾을 수 없습니다:\n{e}")
            return
        except PermissionError:
            messagebox.showerror("권한 오류", "이 파일을 쓸 권한이 없어요. 다른 위치로 저장해봐!")
            return
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("예상치 못한 오류", f"{e}")
            return

        # 4) 성공 안내 + 열기
        messagebox.showinfo("저장 완료", f"PDF로 저장했어:\n{pdf_path}")
        try:
            open_file(pdf_path)
        except Exception:
            # 열기 실패해도 최소한 폴더는 열어주기
            try:
                import os
                folder = os.path.dirname(pdf_path)
                if folder:
                    os.startfile(folder)
            except Exception:
                pass
    tk.Button(root, text="이미지 추가", command=add_image_to_pdf).pack(pady=6)
    tk.Button(root, text="PDF로 저장", command=save_edited_pdf).pack(pady=10)
